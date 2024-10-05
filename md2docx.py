import sys
import os
from typing import Tuple

def check_dependencies() -> bool:
    try:
        import markdown
        import docx
        from bs4 import BeautifulSoup
        return True
    except ImportError:
        return False

def install_dependencies() -> None:
    print("Installing required dependencies...")
    os.system(f"{sys.executable} -m pip install markdown python-docx beautifulsoup4")

def get_file_paths() -> Tuple[str, str]:
    input_file = input("Enter the path to the input Markdown file: ").strip()
    output_file = input("Enter the path for the output Word document: ").strip()
    
    if not output_file.endswith('.docx'):
        output_file += '.docx'
    
    return input_file, output_file

def convert_md_to_docx(input_file: str, output_file: str) -> None:
    import markdown
    from docx import Document
    from docx.shared import Pt
    from bs4 import BeautifulSoup

    # Read the Markdown content
    with open(input_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # Convert Markdown to HTML
    html_content = markdown.markdown(md_content)
    
    # Parse HTML content
    soup = BeautifulSoup(html_content, 'html.parser')
    
    # Create a new Word document
    doc = Document()
    
    # Modify existing styles
    style_h1 = doc.styles['Heading 1']
    style_h1.font.size = Pt(18)
    style_h1.font.bold = True
    
    style_h2 = doc.styles['Heading 2']
    style_h2.font.size = Pt(16)
    style_h2.font.bold = True
    
    style_normal = doc.styles['Normal']
    style_normal.font.size = Pt(11)

    # Convert HTML to Word document
    for element in soup.find_all():
        if element.name == 'h1':
            doc.add_paragraph(element.text, style='Heading 1')
        elif element.name == 'h2':
            doc.add_paragraph(element.text, style='Heading 2')
        elif element.name == 'p':
            doc.add_paragraph(element.text, style='Normal')
        elif element.name == 'ul':
            for li in element.find_all('li'):
                doc.add_paragraph(li.text, style='List Bullet')
        elif element.name == 'ol':
            for li in element.find_all('li'):
                doc.add_paragraph(li.text, style='List Number')
    
    # Save the document
    doc.save(output_file)

def main() -> None:
    if not check_dependencies():
        print("Required dependencies are missing.")
        install = input("Do you want to install them? (y/n): ").strip().lower()
        if install == 'y':
            install_dependencies()
        else:
            print("Cannot proceed without required dependencies. Exiting.")
            sys.exit(1)
    
    input_file, output_file = get_file_paths()
    
    try:
        convert_md_to_docx(input_file, output_file)
        print(f"Conversion complete. Output saved to {output_file}")
    except FileNotFoundError:
        print(f"Error: The input file '{input_file}' was not found.")
    except PermissionError:
        print(f"Error: Permission denied when trying to read '{input_file}' or write '{output_file}'.")
    except Exception as e:
        print(f"An unexpected error occurred: {str(e)}")

if __name__ == "__main__":
    main()